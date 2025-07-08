from django.conf import settings
from django.shortcuts import render,HttpResponse
from numpy import ptp

# Create your views here.
def index(request):
    return render(request,'index.html')
    # return HttpResponse("this is my final project.")
def generate_question(request):
    return render(request,'page1.html')
def Evaluator(request):
    return render(request,'page2.html')
# generator/views.py
import os
import logging
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import UploadedTextForm
from django.core.files.storage import default_storage
from docx import Document
from .generate_questions import read_docx, generate_questions
from docx import Document
from pdf2docx import Converter

from django.shortcuts import render
from django.core.files.storage import default_storage
import os
from django.conf import settings
from django.core.files.storage import default_storage
from django.shortcuts import render
from pdf2docx import Converter
from docx.shared import Pt
from docx import Document
import random
# Configure logging
logging.basicConfig(level=logging.DEBUG)


def upload_text(request):
    if request.method == 'POST':
        logging.debug("Form is being submitted")
        form = UploadedTextForm(request.POST, request.FILES)
        if form.is_valid():
            logging.debug("Form is valid")
            # Get form data
            title = request.POST['title']
            class_name = request.POST['class_name']
            subject = request.POST['subject']
            exam_name = request.POST['exam_name']
            total_questions = int(request.POST['total_questions'])
            duration = request.POST['duration']
            full_marks = request.POST['full_marks']
            
            # Process uploaded file
            uploaded_file = request.FILES['file']
            file_path = os.path.join(settings.MEDIA_ROOT, uploaded_file.name)
            with open(file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            logging.debug(f"File saved to {file_path}")
            
            try:
                # Extract text from DOCX
                text_content = read_docx(file_path)
            except Exception as e:
                logging.error(f"Error reading DOCX file: {e}")
                return HttpResponse(f"Error reading DOCX file: {e}", status=500)
            
            # Generate questions using AI model
            questions = generate_questions(text_content, num_questions=total_questions, question_type="short")
            
            # Create the generated_docs directory if it doesn't exist
            generated_docs_dir = os.path.join(settings.MEDIA_ROOT, 'generated_docs')
            if not os.path.exists(generated_docs_dir):
                os.makedirs(generated_docs_dir)

            # Creating the doc file format
            new_document = Document()
            new_document.add_heading(f'Exam: {exam_name}', 0)
            new_document.add_heading(f'Subject: {subject}', level=1)
            new_document.add_heading(f'Class: {class_name}', level=2)
            new_document.add_heading(f'Duration: {duration}', level=3)
            new_document.add_heading(f'Full Marks: {full_marks}', level=3)
            new_document.add_heading('Questions:', level=1)

            for i, question in enumerate(questions, start=1):
                new_document.add_heading(f'Question {i}:', level=2)
                new_document.add_paragraph(question)

                # Add a text box for writing the answer
                text_box = new_document.add_paragraph()
                text_box.add_run('Answer: ____________________________________________').font.size = Pt(12)

            # Save the document
            generated_file_path = os.path.join(generated_docs_dir, 'exam_question.docx')
            new_document.save(generated_file_path)
            
            # Serve the document for download
            with open(generated_file_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename="exam_question.docx"'
                
                # Close the file before attempting to remove it
                f.close()
                
                # Clean up the saved files
                if os.path.exists(file_path):
                    os.remove(file_path)
                if os.path.exists(generated_file_path):
                    os.remove(generated_file_path)
                
                return response
        else:
            logging.debug("Form is invalid")
    else:
        form = UploadedTextForm()
    
    return render(request, 'page1.html', {'form': form})





from docx import Document
import random

def evaluate_answer(request):
    if request.method == 'POST':
        # Simulated extraction of exam details from the uploaded file
        exam_details = {
            'exam_name': 'Mathematics Mid-Term',
            'subject': 'Mathematics',
            'class': '10th Grade',
            'duration': '120 minutes',
            'full_marks': 100,
        }
        
        # Simulated marks calculation
        total_marks = exam_details['full_marks']
        obtained_marks = random.randint(40, total_marks)  # Simulated marks obtained
        
        # Generate marksheet
        marksheet = Document()
        marksheet.add_heading('MID-TERM EXAMINATION MARKSHEET', 0)
        marksheet.add_heading(f'{exam_details["subject"]} - {exam_details["class"]}', level=1)
        marksheet.add_paragraph(f'Exam Name: {exam_details["exam_name"]}')
        marksheet.add_paragraph(f'Duration: {exam_details["duration"]}')
        marksheet.add_paragraph(f'Full Marks: {total_marks}')
        marksheet.add_paragraph(f'Marks Obtained: {obtained_marks}')
        
        # Styling the marksheet
        section = marksheet.sections[0]
        section.top_margin = Pt(30)
        section.bottom_margin = Pt(30)
        section.left_margin = Pt(60)
        section.right_margin = Pt(60)
        
        style = marksheet.styles['Heading 1']
        font = style.font
        font.bold = True
        font.size = Pt(14)
        
        # Save the marksheet
        marksheet_filename = 'marksheet.docx'
        marksheet.save(marksheet_filename)
        
        # Serve the marksheet for download
        with open(marksheet_filename, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{marksheet_filename}"'
            return response
    
    # Render the page for uploading the file
    return render(request, 'page2.html')
